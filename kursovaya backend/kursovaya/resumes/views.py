# resumes/views.py
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth import login, authenticate
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse
from django.db.models import Q
from docx import Document
import logging
from .forms import UserRegistrationForm, ResumeForm, LoginForm
from .models import Resume
from django.contrib.auth import logout as auth_logout
from django.views.decorators.http import require_POST

logger = logging.getLogger(__name__)
def index(request):
    if request.user.is_authenticated:
        return redirect('home')
    return redirect('register')
@login_required
def resume_list(request):
    resumes = Resume.objects.filter(
        Q(is_public=True) | Q(user=request.user)
    ).distinct()
    return render(request, 'resumes/resume_list.html', {'resumes': resumes})
def resume_detail(request, pk):
    try:
        resume = get_object_or_404(Resume, pk=pk)
        if not resume.is_public and (not request.user.is_authenticated or resume.user != request.user):
            messages.error(request, "У вас нет доступа к этому резюме.")
            return redirect('resume_list')
    except Resume.DoesNotExist:
        messages.error(request, "Резюме не найдено.")
        return redirect('resume_list')
    return render(request, 'resumes/resume_detail.html', {'resume': resume})
def index(request):
    if request.user.is_authenticated:
        return redirect('home')
    return redirect('register')
@login_required
def resume_form(request, pk=None):
    if pk:
        resume = get_object_or_404(Resume, pk=pk, user=request.user)
        if request.method == 'POST':
            form = ResumeForm(request.POST, instance=resume)
            if form.is_valid():
                try:
                    form.save()
                    messages.success(request, "Резюме успешно обновлено!")
                    return redirect('resume_detail', pk=resume.pk)
                except Exception as e:
                    logger.error("Ошибка при обновлении резюме: %s", str(e))
                    messages.error(request, "Ошибка при сохранении.")
            else:
                logger.warning("Ошибки формы при редактировании: %s", form.errors)
        else:
            form = ResumeForm(instance=resume)
        title = 'Редактировать резюме'
    else:
        if request.method == 'POST':
            form = ResumeForm(request.POST)
            if form.is_valid():
                try:
                    resume = form.save(commit=False)
                    resume.user = request.user
                    resume.is_public = True
                    resume.save()
                    messages.success(request, "Резюме успешно создано!")
                    return redirect('resume_detail', pk=resume.pk)
                except Exception as e:
                    logger.error("Ошибка при создании резюме: %s", str(e))
                    messages.error(request, "Произошла ошибка при сохранении.")
            else:
                logger.warning("Ошибки формы при создании: %s", form.errors)
        else:
            form = ResumeForm()
        title = 'Создать резюме'

    return render(request, 'resumes/resume_form.html', {
        'form': form,
        'title': title,
        'pk': pk
    })
@login_required
def resume_delete(request, pk):
    try:
        resume = get_object_or_404(Resume, pk=pk)
        if resume.user != request.user and not resume.is_public:
            messages.error(request, "У вас нет прав для удаления этого резюме.")
            return redirect('resume_list')
        resume.delete()
        messages.success(request, "Резюме успешно удалено.")
        return redirect('resume_list')
    except Resume.DoesNotExist:
        messages.error(request, "Резюме не найдено.")
        return redirect('resume_list')

@login_required
def export_resume_docx(request, pk):
    try:
        resume = get_object_or_404(Resume, pk=pk)
        if not resume.is_public and resume.user != request.user:
            messages.error(request, "У вас нет доступа к экспорту этого резюме.")
            return redirect('resume_list')
        document = Document()
        document.add_heading(resume.full_name or "Без имени", 0)
        if resume.email or resume.phone or resume.date_of_birth or resume.address:
            document.add_heading("Личная информация", level=1)
            if resume.email:
                document.add_paragraph(f"Email: {resume.email}")
            if resume.phone:
                document.add_paragraph(f"Телефон: {resume.phone}")
            if resume.date_of_birth:
                document.add_paragraph(f"Дата рождения: {resume.date_of_birth.strftime('%d.%m.%Y')}")
            if resume.address:
                document.add_paragraph(f"Адрес: {resume.address}")
        if resume.summary or resume.skills or resume.experience or resume.education or resume.hobby or resume.project:
            document.add_heading("Дополнительная информация", level=1)
            if resume.summary:
                document.add_paragraph(f"Описание: {resume.summary}")
            if resume.skills:
                document.add_paragraph(f"Навыки: {resume.skills}")
            if resume.experience:
                document.add_paragraph(f"Опыт работы: {resume.experience}")
            if resume.education:
                document.add_paragraph(f"Образование: {resume.education}")
            if resume.hobby:
                document.add_paragraph(f"Хобби: {resume.hobby}")
            if resume.project:
                document.add_paragraph(f"Проект: {resume.project}")
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename=resume_{resume.pk}_{resume.full_name or 'unnamed'}.docx'
        document.save(response)
        return response
    except Resume.DoesNotExist:
        messages.error(request, "Резюме не найдено.")
        return redirect('resume_list')
@login_required
def home_view(request):
    user = request.user
    resumes = Resume.objects.filter(user=user).order_by('-created_at')
    recent_resumes = resumes[:4]
    total_resumes = resumes.count()
    progress = 0
    missing_fields = []
    if resumes.exists():
        last = resumes[0]
        field_map = {
            'ФИО': last.full_name,
            'Дата рождения': last.date_of_birth,
            'Email': last.email,
            'Телефон': last.phone,
            'Адрес': last.address,
            'О себе': last.summary,
            'Навыки': last.skills,
            'Опыт': last.experience,
            'Образование': last.education,
        }
        filled = sum(bool(value) for value in field_map.values())
        progress = int((filled / len(field_map)) * 100)
        missing_fields = [name for name, value in field_map.items() if not value]

    return render(request, 'resumes/home.html', {
        'recent_resumes': recent_resumes,
        'total_resumes': total_resumes,
        'progress': progress,
        'missing_fields': missing_fields,
    })

def register(request):
    if request.method == 'POST':
        form = UserRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('home')
    else:
        form = UserRegistrationForm()
    return render(request, 'registration/register.html', {'form': form})
@login_required
def user_profile_view(request):
    return render(request, 'accounts/profile.html')
def custom_logout(request):
    logger.info(f"Logout request method: {request.method}")
    auth_logout(request)
    return redirect('login')  # Перенаправляем на регистрацию после выхода
def user_login(request):
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data['username']
            password = form.cleaned_data['password']
            user = authenticate(request, username=username, password=password)

            if user:
                login(request, user)
                return redirect('home')
            else:
                error = 'Неверный логин или пароль'
    else:
        form = LoginForm()

    return render(request, 'registration/login.html', context={'form': form})
